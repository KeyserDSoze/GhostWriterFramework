#!/usr/bin/env python3
"""
Narrarium Manuscript Builder

Generates submission-ready .docx manuscripts from a Narrarium book repository.
Follows standard manuscript format conventions:
  - Times New Roman 12pt (configurable)
  - Double-spaced
  - 1-inch margins
  - Title page with title, author, word count
  - Running header: Author / Title / Page
  - 0.5-inch first-line paragraph indent
  - Chapter breaks on new pages
  - Scene breaks marked with #

Outputs:
  - Full manuscript: build/manuscript.docx
  - Sample chapters: build/manuscript-sample.docx (first N chapters, default 5)

Usage:
  python build_manuscript.py [--book-root <path>] [--config <path>]

Reads settings from manuscript.yaml in the book root.
"""

import argparse
import os
import re
import sys
from pathlib import Path

try:
    import yaml
except ImportError:
    yaml = None

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print(
        "ERROR: python-docx is required. Install it with:\n"
        "  pip install python-docx\n"
        "Or:\n"
        "  pip install -r requirements-manuscript.txt",
        file=sys.stderr,
    )
    sys.exit(1)

# ---------------------------------------------------------------------------
# Frontmatter parsing
# ---------------------------------------------------------------------------

_FRONTMATTER_RE = re.compile(r"^---\s*\n(.*?)\n---\s*\n?", re.DOTALL)


def strip_frontmatter(text: str) -> tuple[dict, str]:
    """Return (frontmatter_dict, body) from a markdown file with YAML front matter."""
    match = _FRONTMATTER_RE.match(text)
    if not match:
        return {}, text

    fm_raw = match.group(1)
    body = text[match.end():]

    fm: dict = {}
    if yaml is not None:
        try:
            fm = yaml.safe_load(fm_raw) or {}
        except Exception:
            fm = {}
    else:
        # Minimal key: value parser when PyYAML is not available
        for line in fm_raw.splitlines():
            if ":" in line:
                key, _, value = line.partition(":")
                key = key.strip()
                value = value.strip().strip('"').strip("'")
                if value:
                    fm[key] = value

    return fm, body


# ---------------------------------------------------------------------------
# Book reading helpers
# ---------------------------------------------------------------------------


def read_book_metadata(book_root: Path) -> dict:
    """Read book.md and return its frontmatter."""
    book_file = book_root / "book.md"
    if not book_file.exists():
        print(f"ERROR: {book_file} not found. Are you in a Narrarium book repo?", file=sys.stderr)
        sys.exit(1)
    fm, _ = strip_frontmatter(book_file.read_text(encoding="utf-8"))
    return fm


def list_chapters(book_root: Path) -> list[dict]:
    """Return sorted list of chapter dicts with slug, metadata, and path."""
    chapters_dir = book_root / "chapters"
    if not chapters_dir.is_dir():
        return []

    chapters = []
    for entry in sorted(chapters_dir.iterdir()):
        if not entry.is_dir():
            continue
        chapter_file = entry / "chapter.md"
        if not chapter_file.exists():
            continue
        fm, body = strip_frontmatter(chapter_file.read_text(encoding="utf-8"))
        chapters.append({
            "slug": entry.name,
            "path": entry,
            "metadata": fm,
            "body": body,
            "number": fm.get("number", 0),
        })

    chapters.sort(key=lambda c: c["number"])
    return chapters


def read_paragraphs(chapter_path: Path) -> list[dict]:
    """Return sorted list of paragraph dicts for a chapter directory."""
    skip_files = {"chapter.md", "writing-style.md", "notes.md", "ideas.md", "promoted.md"}
    paragraphs = []

    for md_file in sorted(chapter_path.glob("*.md")):
        if md_file.name in skip_files:
            continue
        fm, body = strip_frontmatter(md_file.read_text(encoding="utf-8"))
        paragraphs.append({
            "metadata": fm,
            "body": body.strip(),
            "number": fm.get("number", 0),
        })

    paragraphs.sort(key=lambda p: p["number"])
    return paragraphs


# ---------------------------------------------------------------------------
# Settings
# ---------------------------------------------------------------------------

DEFAULT_SETTINGS = {
    "sample_chapters": 5,
    "show_paragraph_titles": False,
    "show_chapter_summary": False,
    "font_name": "Times New Roman",
    "font_size": 12,
    "line_spacing": 2.0,
    "margin_inches": 1.0,
    "paragraph_indent_inches": 0.5,
    "paragraph_break_newlines": 3,
    "scene_break": "#",
    "output_dir": "build",
    "full_filename": "manuscript.docx",
    "sample_filename": "manuscript-sample.docx",
    "include_title_page": True,
    "page_size": "letter",
}


def load_settings(book_root: Path, config_path: str | None = None) -> dict:
    """Load manuscript.yaml settings, merged with defaults."""
    settings = dict(DEFAULT_SETTINGS)

    candidates = []
    if config_path:
        candidates.append(Path(config_path))
    candidates.append(book_root / "manuscript.yaml")
    candidates.append(book_root / "manuscript.yml")

    for candidate in candidates:
        if candidate.exists():
            raw = candidate.read_text(encoding="utf-8")
            if yaml is not None:
                user = yaml.safe_load(raw) or {}
            else:
                user = {}
                for line in raw.splitlines():
                    if ":" in line and not line.strip().startswith("#"):
                        key, _, value = line.partition(":")
                        key = key.strip()
                        value = value.strip().strip('"').strip("'")
                        if value.lower() == "true":
                            value = True
                        elif value.lower() == "false":
                            value = False
                        else:
                            try:
                                value = int(value)
                            except ValueError:
                                try:
                                    value = float(value)
                                except ValueError:
                                    pass
                        if value != "":
                            user[key] = value
            settings.update(user)
            break

    return settings


# ---------------------------------------------------------------------------
# Word count helper
# ---------------------------------------------------------------------------


def count_words(text: str) -> int:
    """Simple word count on plain text."""
    return len(text.split())


# ---------------------------------------------------------------------------
# DOCX builder
# ---------------------------------------------------------------------------


def set_page_size_and_margins(section, settings: dict):
    """Configure page size and margins."""
    if settings["page_size"] == "a4":
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    else:  # letter
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    margin = Inches(settings["margin_inches"])
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin


def add_header(section, author: str, title: str, font_name: str, font_size: int):
    """Add running header: Author / Title / Page Number."""
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    short_title = title[:30].upper() if len(title) > 30 else title.upper()
    author_last = author.split()[-1] if author else "AUTHOR"
    run = paragraph.add_run(f"{author_last} / {short_title} / ")
    run.font.name = font_name
    run.font.size = Pt(font_size)

    # Page number field
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run_page = paragraph.add_run()
    run_page._r.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    run_page2 = paragraph.add_run()
    run_page2.font.name = font_name
    run_page2.font.size = Pt(font_size)
    run_page2._r.append(instr_text)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_page3 = paragraph.add_run()
    run_page3._r.append(fld_char_end)


def configure_style(doc: "Document", settings: dict):
    """Set up the default body style and heading styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = settings["font_name"]
    font.size = Pt(settings["font_size"])
    pf = style.paragraph_format
    pf.line_spacing = settings["line_spacing"]
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Inches(settings["paragraph_indent_inches"])

    # Heading 1 for chapter titles
    if "Heading 1" in doc.styles:
        h1 = doc.styles["Heading 1"]
        h1.font.name = settings["font_name"]
        h1.font.size = Pt(settings["font_size"])
        h1.font.bold = True
        h1.font.color.rgb = None  # black
        h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1.paragraph_format.space_before = Pt(72)  # ~1 inch from top
        h1.paragraph_format.space_after = Pt(24)
        h1.paragraph_format.first_line_indent = Inches(0)
        h1.paragraph_format.line_spacing = settings["line_spacing"]
        h1.paragraph_format.page_break_before = True

    # Heading 2 for paragraph/scene titles (optional)
    if "Heading 2" in doc.styles:
        h2 = doc.styles["Heading 2"]
        h2.font.name = settings["font_name"]
        h2.font.size = Pt(settings["font_size"])
        h2.font.bold = True
        h2.font.italic = False
        h2.font.color.rgb = None
        h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h2.paragraph_format.space_before = Pt(12)
        h2.paragraph_format.space_after = Pt(6)
        h2.paragraph_format.first_line_indent = Inches(0)
        h2.paragraph_format.line_spacing = settings["line_spacing"]


def add_title_page(doc: "Document", book: dict, word_count: int, settings: dict):
    """Add a standard manuscript title page."""
    # Contact info block (top-left)
    author = book.get("author", "")
    if author:
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.LEFT
        contact.paragraph_format.first_line_indent = Inches(0)
        contact.paragraph_format.space_after = Pt(0)
        run = contact.add_run(author)
        run.font.name = settings["font_name"]
        run.font.size = Pt(settings["font_size"])

    # Word count (top-right aligned paragraph)
    wc_para = doc.add_paragraph()
    wc_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    wc_para.paragraph_format.first_line_indent = Inches(0)
    wc_para.paragraph_format.space_after = Pt(0)
    rounded_wc = round(word_count / 100) * 100 if word_count > 500 else word_count
    run = wc_para.add_run(f"Approx. {rounded_wc:,} words")
    run.font.name = settings["font_name"]
    run.font.size = Pt(settings["font_size"])

    # Vertical space to center the title
    for _ in range(10):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.first_line_indent = Inches(0)
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(0)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.first_line_indent = Inches(0)
    title_para.paragraph_format.space_after = Pt(12)
    run = title_para.add_run(book.get("title", "Untitled"))
    run.font.name = settings["font_name"]
    run.font.size = Pt(settings["font_size"])
    run.bold = True

    # "by" line
    if author:
        by_para = doc.add_paragraph()
        by_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        by_para.paragraph_format.first_line_indent = Inches(0)
        run = by_para.add_run(f"by {author}")
        run.font.name = settings["font_name"]
        run.font.size = Pt(settings["font_size"])

    # Genre line if available
    genre = book.get("genre")
    if genre:
        genre_para = doc.add_paragraph()
        genre_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        genre_para.paragraph_format.first_line_indent = Inches(0)
        run = genre_para.add_run(genre)
        run.font.name = settings["font_name"]
        run.font.size = Pt(settings["font_size"])


def add_custom_frontmatter_fields(doc: "Document", metadata: dict, settings: dict):
    """Render show_custom_* fields from settings, looking up values in metadata.

    Settings like ``show_custom_date: 11`` and ``show_custom_fancyname: 10``
    will look for ``date`` and ``fancyname`` in *metadata*, sort by numeric
    order (10 before 11), and add an italic paragraph for each found value.
    """
    customs: list[tuple[int, str]] = []
    for key, order in settings.items():
        if key.startswith("show_custom_") and isinstance(order, (int, float)):
            field_name = key[len("show_custom_"):]
            customs.append((int(order), field_name))
    if not customs:
        return
    customs.sort(key=lambda t: t[0])
    for _, field_name in customs:
        value = metadata.get(field_name)
        if not value:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(str(value))
        run.font.name = settings["font_name"]
        run.font.size = Pt(settings["font_size"])
        run.italic = True


def add_chapter(doc: "Document", chapter: dict, settings: dict, is_first_chapter: bool = False):
    """Add a chapter with all its paragraphs to the document."""
    chapter_meta = chapter["metadata"]
    chapter_title = chapter_meta.get("title", f"Chapter {chapter_meta.get('number', '?')}")
    chapter_number = chapter_meta.get("number", "")

    # Chapter heading (page break is built into the Heading 1 style)
    heading_text = f"Chapter {chapter_number}" if chapter_number else chapter_title
    h = doc.add_heading(heading_text, level=1)

    # If first chapter after title page, suppress the extra page break
    # (the title page already ends on its own page)
    if is_first_chapter:
        h.paragraph_format.page_break_before = True

    # Chapter title under the number if different
    if chapter_number and chapter_title != f"Chapter {chapter_number}":
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.first_line_indent = Inches(0)
        subtitle.paragraph_format.space_after = Pt(24)
        run = subtitle.add_run(chapter_title)
        run.font.name = settings["font_name"]
        run.font.size = Pt(settings["font_size"])
        run.italic = True

    # Chapter summary if enabled
    if settings["show_chapter_summary"]:
        summary = chapter_meta.get("summary", "")
        if summary:
            summary_para = doc.add_paragraph()
            summary_para.paragraph_format.first_line_indent = Inches(0)
            summary_para.paragraph_format.space_after = Pt(12)
            run = summary_para.add_run(summary)
            run.font.name = settings["font_name"]
            run.font.size = Pt(settings["font_size"])
            run.italic = True

    # Custom frontmatter fields from chapter metadata (show_custom_*)
    add_custom_frontmatter_fields(doc, chapter_meta, settings)

    paragraphs = read_paragraphs(chapter["path"])

    for idx, para in enumerate(paragraphs):
        # Paragraph/scene title if enabled
        if settings["show_paragraph_titles"]:
            para_title = para["metadata"].get("title", "")
            if para_title:
                h2 = doc.add_heading(para_title, level=2)

        # Scene break between paragraphs (not before the first one)
        if idx > 0 and not settings["show_paragraph_titles"]:
            scene_break = doc.add_paragraph()
            scene_break.alignment = WD_ALIGN_PARAGRAPH.CENTER
            scene_break.paragraph_format.first_line_indent = Inches(0)
            scene_break.paragraph_format.space_before = Pt(12)
            scene_break.paragraph_format.space_after = Pt(12)
            run = scene_break.add_run(settings["scene_break"])
            run.font.name = settings["font_name"]
            run.font.size = Pt(settings["font_size"])

        # Custom frontmatter fields from paragraph metadata (show_custom_*)
        add_custom_frontmatter_fields(doc, para["metadata"], settings)

        # Body text
        body = para["body"]
        if not body:
            continue

        # Split body into paragraphs using the configured newline threshold.
        break_nl = settings.get("paragraph_break_newlines", 3)
        sep = "\n" * max(break_nl, 2)
        for text_para in body.split(sep):
            text_para = text_para.strip()
            if not text_para:
                continue
            # Collapse remaining newlines (below threshold) into spaces.
            text_para = re.sub(r"\n+", " ", text_para)
            p = doc.add_paragraph(text_para)
            # The Normal style already has the right formatting


def build_manuscript(
    book_root: Path,
    chapters: list[dict],
    book_meta: dict,
    settings: dict,
    output_path: Path,
    label: str = "manuscript",
):
    """Build a single .docx manuscript file."""
    doc = Document()

    # Set compatibility mode to Word 2016+ (version 15) to avoid
    # "older file type" warnings in modern Word.
    import lxml.etree as ET
    compat = doc.settings.element.makeelement(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}compat", {}
    )
    compat_setting = compat.makeelement(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}compatSetting",
        {
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name": "compatibilityMode",
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}uri": "http://schemas.microsoft.com/office/word",
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val": "15",
        },
    )
    compat.append(compat_setting)
    doc.settings.element.append(compat)

    configure_style(doc, settings)

    section = doc.sections[0]
    set_page_size_and_margins(section, settings)

    # Calculate total word count
    total_words = 0
    for ch in chapters:
        for para in read_paragraphs(ch["path"]):
            total_words += count_words(para["body"])

    # Running header
    add_header(section, book_meta.get("author", ""), book_meta.get("title", ""), settings["font_name"], settings["font_size"])

    # Title page
    if settings["include_title_page"]:
        add_title_page(doc, book_meta, total_words, settings)

    # Chapters
    for idx, chapter in enumerate(chapters):
        add_chapter(doc, chapter, settings, is_first_chapter=(idx == 0))

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    return total_words


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main():
    parser = argparse.ArgumentParser(
        description="Build submission-ready .docx manuscripts from a Narrarium book.",
    )
    parser.add_argument(
        "--book-root",
        default=".",
        help="Path to the Narrarium book root (default: current directory)",
    )
    parser.add_argument(
        "--config",
        default=None,
        help="Path to manuscript.yaml config file (default: <book-root>/manuscript.yaml)",
    )
    parser.add_argument(
        "--full-only",
        action="store_true",
        help="Only generate the full manuscript, skip the sample",
    )
    parser.add_argument(
        "--sample-only",
        action="store_true",
        help="Only generate the sample manuscript",
    )
    args = parser.parse_args()

    book_root = Path(args.book_root).resolve()
    settings = load_settings(book_root, args.config)
    book_meta = read_book_metadata(book_root)
    chapters = list_chapters(book_root)

    if not chapters:
        print("WARNING: No chapters found. Nothing to build.", file=sys.stderr)
        sys.exit(0)

    output_dir = book_root / settings["output_dir"]
    results = []

    # Full manuscript
    if not args.sample_only:
        full_path = output_dir / settings["full_filename"]
        word_count = build_manuscript(book_root, chapters, book_meta, settings, full_path, "full manuscript")
        results.append(("Full manuscript", full_path, len(chapters), word_count))

    # Sample manuscript
    if not args.full_only:
        sample_count = int(settings["sample_chapters"])
        if sample_count > 0 and len(chapters) >= sample_count:
            sample_chapters = chapters[:sample_count]
            sample_path = output_dir / settings["sample_filename"]
            sample_words = build_manuscript(book_root, sample_chapters, book_meta, settings, sample_path, "sample")
            results.append(("Sample manuscript", sample_path, len(sample_chapters), sample_words))
        elif sample_count > 0 and len(chapters) < sample_count:
            print(
                f"NOTE: Book has {len(chapters)} chapters, fewer than sample_chapters={sample_count}. "
                "Skipping sample generation (full manuscript already covers all chapters).",
                file=sys.stderr,
            )

    # Summary
    print()
    print("=" * 60)
    print("  NARRARIUM MANUSCRIPT BUILD COMPLETE")
    print("=" * 60)
    for label, path_val, ch_count, wc in results:
        print(f"  {label}:")
        print(f"    File:     {path_val}")
        print(f"    Chapters: {ch_count}")
        print(f"    Words:    ~{wc:,}")
        print()
    print("  Format: Standard Manuscript Format")
    print(f"  Font:   {settings['font_name']} {settings['font_size']}pt")
    print(f"  Spacing: {settings['line_spacing']}x (double-spaced)")
    print(f"  Margins: {settings['margin_inches']}\"")
    print("=" * 60)


if __name__ == "__main__":
    main()
